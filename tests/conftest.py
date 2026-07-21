"""Shared fixtures: sample docx/pdf files, fake LLMs, tmp data dir."""

import io

import pytest
from docx import Document

from src import config
from src.models.schemas import (
    CareerProfile,
    Conflict,
    Experience,
    Project,
    Skill,
)


class RawMessage:
    """Minimal stand-in for an AIMessage carrying structured-output tool calls."""

    def __init__(self, args: dict | None = None, name: str = "SourceExtraction"):
        self.tool_calls = [] if args is None else [{"name": name, "args": args}]


class FakeLLM:
    """Stands in for ChatAnthropic().with_structured_output(...).

    `responses` may be a single object (returned every call), a list
    (returned in order), or a callable(messages) -> object.

    With `include_raw=True` the response is wrapped in the
    `{"parsed", "raw", "parsing_error"}` envelope LangChain returns. A response
    that is already such a dict (used to simulate a validation failure) is
    passed through untouched.
    """

    def __init__(self, responses):
        self.responses = responses
        self.calls: list = []
        self.include_raw = False

    def with_structured_output(self, schema, include_raw: bool = False):
        self.include_raw = include_raw
        return self

    def invoke(self, messages):
        self.calls.append(messages)
        if callable(self.responses):
            response = self.responses(messages)
        elif isinstance(self.responses, list):
            response = self.responses.pop(0)
        else:
            response = self.responses
        if not self.include_raw or isinstance(response, dict):
            return response
        return {
            "parsed": response,
            "raw": RawMessage(response.model_dump()),
            "parsing_error": None,
        }


@pytest.fixture
def data_dir(tmp_path, monkeypatch):
    """Point the profile store at a temporary directory."""
    monkeypatch.setattr(config, "DATA_DIR", tmp_path)
    return tmp_path


def build_sample_docx() -> bytes:
    doc = Document()
    doc.add_heading("Alice Smith", level=0)
    doc.add_paragraph("alice@example.com | London")
    doc.add_heading("Experience", level=1)
    doc.add_paragraph("Senior Engineer, Acme Corp, 2020-2024")
    doc.add_paragraph("Built a distributed trading backtester in Python")
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, PostgreSQL, Docker")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_minimal_pdf(lines: list[str]) -> bytes:
    """Hand-roll a minimal single-page PDF that pdfplumber can read."""
    parts = ["BT /F1 12 Tf 72 720 Td 14 TL"]
    for i, line in enumerate(lines):
        esc = line.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
        parts.append(f"({esc}) Tj" if i == 0 else f"T* ({esc}) Tj")
    parts.append("ET")
    stream = "\n".join(parts).encode()
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length %d >>\nstream\n%s\nendstream" % (len(stream), stream),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out += b"%d 0 obj\n" % i + obj + b"\nendobj\n"
    xref_pos = len(out)
    out += b"xref\n0 %d\n0000000000 65535 f \n" % (len(objects) + 1)
    for off in offsets:
        out += b"%010d 00000 n \n" % off
    out += b"trailer\n<< /Size %d /Root 1 0 R >>\nstartxref\n%d\n%%%%EOF" % (
        len(objects) + 1,
        xref_pos,
    )
    return bytes(out)


LINKEDIN_CSVS: dict[str, str] = {
    "Profile.csv": (
        "First Name,Last Name,Headline,Summary,Geo Location\r\n"
        "Alice,Smith,Senior Engineer,Backend and data engineering.,London\r\n"
    ),
    # LinkedIn prefixes several export files with a free-text note before the
    # header row — the parser must skip it rather than assume line 1.
    "Positions.csv": (
        'Notes:\r\n'
        '"This file contains your positions."\r\n'
        "\r\n"
        "Company Name,Title,Description,Location,Started On,Finished On\r\n"
        "Acme Corp,Senior Engineer,Built a distributed trading backtester in "
        "Python,London,Jan 2021,\r\n"
        "Globex,Engineer,Maintained the billing service,Leeds,Mar 2018,Dec 2020\r\n"
    ),
    "Education.csv": (
        "School Name,Start Date,End Date,Notes,Degree Name,Activities\r\n"
        "University of Leeds,2014,2018,,BSc Computer Science,Robotics society\r\n"
    ),
    "Skills.csv": "Name\r\nPython\r\nPostgreSQL\r\n",
    "Certifications.csv": (
        "Name,Url,Authority,Started On,Finished On,License Number\r\n"
        "AWS Solutions Architect,,Amazon Web Services,Feb 2022,,ABC-123\r\n"
    ),
    "Recommendations_Received.csv": (
        "First Name,Last Name,Company,Job Title,Text,Creation Date,Status\r\n"
        "Bob,Jones,Acme Corp,CTO,Alice rebuilt our data pipeline end to end.,"
        "2022-01-05,VISIBLE\r\n"
    ),
}


def build_linkedin_export_zip(members: dict[str, str] | None = None) -> bytes:
    """Build an official-shaped LinkedIn data-export ZIP for tests."""
    import zipfile

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as archive:
        for name, text in (members or LINKEDIN_CSVS).items():
            archive.writestr(name, text)
    return buf.getvalue()


@pytest.fixture
def sample_linkedin_zip(tmp_path):
    path = tmp_path / "Basic_LinkedInDataExport.zip"
    path.write_bytes(build_linkedin_export_zip())
    return path


@pytest.fixture
def sample_docx(tmp_path):
    path = tmp_path / "resume.docx"
    path.write_bytes(build_sample_docx())
    return path


@pytest.fixture
def sample_pdf(tmp_path):
    path = tmp_path / "resume.pdf"
    path.write_bytes(
        build_minimal_pdf(
            [
                "Alice Smith",
                "Senior Engineer, Acme Corp, 2020-2024",
                "Built a distributed trading backtester in Python",
            ]
        )
    )
    return path


@pytest.fixture
def sample_profile() -> CareerProfile:
    from src.agents.synthesis import build_raw_source_map

    profile = CareerProfile(
        name="Alice Smith",
        headline="Senior Engineer",
        contact={"email": "alice@example.com"},
        experiences=[
            Experience(
                company="Acme Corp",
                title="Senior Engineer",
                start_date="2020",
                end_date="2024",
                bullets=[
                    "Built a distributed trading backtester in Python",
                    "Led migration of the data pipeline to PostgreSQL",
                ],
                source="cv_docx:resume.docx",
            )
        ],
        projects=[
            Project(
                name="backtester",
                description="Open-source distributed backtesting engine",
                technologies=["Python"],
                source="github:alice",
            )
        ],
        skills=[
            Skill(name="Python", category="language", evidence_count=2),
            Skill(name="PostgreSQL", category="tool", evidence_count=1),
        ],
        certifications=[],
        summary_narrative="Alice is a senior engineer.",
        conflicts=[
            Conflict(
                field="experience.start_date",
                description="CV and GitHub disagree on start date",
                values={"cv_docx:resume.docx": "2020", "github:alice": "2019"},
            )
        ],
    )
    profile.raw_source_map = build_raw_source_map(profile)
    return profile
