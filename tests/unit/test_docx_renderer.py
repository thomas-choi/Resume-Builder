"""Pure .docx rendering: re-open the output and assert content + ordering."""

from docx import Document

from src import config
from src.models.schemas import CoverLetter, Experience, Project, TailoredCV
from src.tools import docx_renderer


def _tailored_cv() -> TailoredCV:
    return TailoredCV(
        headline="Senior Backend Engineer — Python & data pipelines",
        summary="Eight years building data-heavy backends.",
        selected_experiences=[
            Experience(
                company="Acme Corp",
                title="Senior Engineer",
                start_date="2020",
                end_date="2024",
                location="London",
                bullets=[
                    "Built a distributed trading backtester in Python",
                    "Led migration of the data pipeline to PostgreSQL",
                ],
                source="cv_docx:resume.docx",
            )
        ],
        selected_projects=[
            Project(
                name="backtester",
                description="Open-source distributed backtesting engine",
                technologies=["Python", "Redis"],
                url="https://github.com/alice/backtester",
                source="github:alice",
            )
        ],
        highlighted_skills=["Python", "PostgreSQL"],
        relevance_notes={"backtester": "matches the distributed-systems requirement"},
    )


def _texts(path) -> list[str]:
    return [p.text for p in Document(str(path)).paragraphs if p.text.strip()]


def test_render_cv_writes_every_section_in_schema_order(tmp_path):
    out = docx_renderer.render_cv(
        _tailored_cv(),
        tmp_path / "cv.docx",
        name="Alice Smith",
        contact={"email": "alice@example.com", "location": "London"},
    )
    assert out.exists()
    texts = _texts(out)

    assert texts[0] == "Alice Smith"
    assert texts[1] == "alice@example.com | London"
    for expected in (
        "Senior Backend Engineer — Python & data pipelines",
        "Summary",
        "Eight years building data-heavy backends.",
        "Experience",
        "Senior Engineer — Acme Corp",
        "2020 – 2024 | London",
        "Built a distributed trading backtester in Python",
        "Projects",
        "backtester",
        "Technologies: Python, Redis",
        "Skills",
        "Python, PostgreSQL",
    ):
        assert expected in texts

    # Section ordering: header -> summary -> experience -> projects -> skills
    assert texts.index("Summary") < texts.index("Experience") < texts.index("Projects")
    assert texts.index("Projects") < texts.index("Skills")


def test_render_cv_omits_internal_relevance_notes(tmp_path):
    out = docx_renderer.render_cv(_tailored_cv(), tmp_path / "cv.docx", name="Alice")
    assert "matches the distributed-systems requirement" not in "\n".join(_texts(out))


def test_render_cv_omits_empty_sections(tmp_path):
    cv = TailoredCV(headline="Engineer", summary="Pitch.")
    texts = _texts(docx_renderer.render_cv(cv, tmp_path / "cv.docx", name="Alice"))
    assert "Experience" not in texts
    assert "Projects" not in texts
    assert "Skills" not in texts


def test_render_cover_letter_orders_greeting_body_closing(tmp_path):
    letter = CoverLetter(
        greeting="Dear Hiring Manager,",
        body_paragraphs=["First paragraph.", "Second paragraph."],
        closing="Sincerely,",
    )
    out = docx_renderer.render_cover_letter(
        letter,
        tmp_path / "cover-letter.docx",
        name="Alice Smith",
        contact={"email": "alice@example.com"},
    )
    texts = _texts(out)
    assert texts == [
        "Alice Smith",
        "alice@example.com",
        "Dear Hiring Manager,",
        "First paragraph.",
        "Second paragraph.",
        "Sincerely,",
        "Alice Smith",  # signature under the closing
    ]


def test_render_uses_a_configured_template(tmp_path, monkeypatch):
    # A template supplies styles/theme; content is always appended by the
    # renderer, so anything already in the template survives the render.
    template = Document()
    template.add_paragraph("COMPANY LETTERHEAD")
    template_path = tmp_path / "template.docx"
    template.save(str(template_path))
    monkeypatch.setattr(config, "DOCX_TEMPLATE", template_path)

    out = docx_renderer.render_cv(_tailored_cv(), tmp_path / "cv.docx", name="Alice")
    assert _texts(out)[0] == "COMPANY LETTERHEAD"


def test_missing_template_falls_back_to_the_default(tmp_path, monkeypatch, caplog):
    monkeypatch.setattr(config, "DOCX_TEMPLATE", tmp_path / "absent.docx")
    out = docx_renderer.render_cv(_tailored_cv(), tmp_path / "cv.docx", name="Alice")
    assert out.exists()
    assert "using the default template" in caplog.text


def test_convert_to_pdf_disabled_returns_none(tmp_path, monkeypatch):
    monkeypatch.setattr(config, "RENDER_PDF", False)
    docx = docx_renderer.render_cv(_tailored_cv(), tmp_path / "cv.docx")
    assert docx_renderer.convert_to_pdf(docx) is None


def test_convert_to_pdf_without_libreoffice_warns_and_returns_none(
    tmp_path, monkeypatch, caplog
):
    # The .docx is the guaranteed output; a missing converter must not fail
    # the tailoring run.
    monkeypatch.setattr(config, "RENDER_PDF", True)
    monkeypatch.setattr(config, "LIBREOFFICE_BIN", "definitely-not-installed")
    docx = docx_renderer.render_cv(_tailored_cv(), tmp_path / "cv.docx")
    assert docx_renderer.convert_to_pdf(docx) is None
    assert "skipping PDF" in caplog.text


def test_convert_to_pdf_invokes_libreoffice_and_returns_the_pdf(tmp_path, monkeypatch):
    monkeypatch.setattr(config, "RENDER_PDF", True)
    docx = docx_renderer.render_cv(_tailored_cv(), tmp_path / "cv.docx")
    recorded = {}

    def fake_run(command, **kwargs):
        recorded["command"] = command
        docx.with_suffix(".pdf").write_bytes(b"%PDF-1.4 fake")
        return None

    monkeypatch.setattr(docx_renderer.subprocess, "run", fake_run)
    pdf = docx_renderer.convert_to_pdf(docx)

    assert pdf == docx.with_suffix(".pdf")
    assert "--headless" in recorded["command"]
    assert "--convert-to" in recorded["command"]
    # A throwaway user profile keeps LibreOffice off a possibly read-only HOME.
    assert any(c.startswith("-env:UserInstallation=") for c in recorded["command"])
