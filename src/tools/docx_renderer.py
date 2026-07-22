"""Deterministic .docx rendering of tailored documents, plus PDF conversion.

Design doc §9 (Document Agent): rendering is a **pure** step — no LLM sees this
module. Everything written here already exists in the `TailoredCV` /
`CoverLetter` the tailoring agent produced and the validation gate checked, so
the renderer never adds, rewrites or infers content; it only lays it out.

Layout is built with `python-docx` on top of an optional base template
(`DOCX_TEMPLATE`), which supplies the styles/theme — content is always appended
by this module, so a template only needs the usual `Heading 1` / `List Bullet`
styles, and a template missing them degrades to bold/plain paragraphs rather
than failing the render.

PDF is a second step: the rendered .docx is converted by headless LibreOffice
(`LIBREOFFICE_BIN`), which is present in the Docker image. Locally it is often
absent, so conversion failures are logged and return ``None`` — the .docx is
the guaranteed output.
"""

import logging
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Pt

from src import config
from src.models.schemas import CoverLetter, Experience, Project, TailoredCV

logger = logging.getLogger(__name__)


def _new_document() -> Document:
    """Open the configured base template, or python-docx's default template."""
    template = config.DOCX_TEMPLATE
    if template and Path(template).exists():
        return Document(str(template))
    if template:
        logger.warning(
            "docx_renderer: DOCX_TEMPLATE %s not found — using the default template",
            template,
        )
    return Document()


def _heading(document: Document, text: str, level: int) -> None:
    """Add a heading, falling back to a bold paragraph if the style is absent."""
    try:
        document.add_heading(text, level=level)
    except KeyError:  # template without the built-in Heading styles
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(16 if level == 0 else 13)


def _bullet(document: Document, text: str) -> None:
    """Add a bulleted line, falling back to a dash-prefixed paragraph."""
    try:
        document.add_paragraph(text, style="List Bullet")
    except KeyError:  # template without the built-in List Bullet style
        document.add_paragraph(f"• {text}")


def _contact_line(contact: dict[str, str] | None) -> str:
    """Join the contact values into one header line, preserving profile order."""
    return " | ".join(str(v).strip() for v in (contact or {}).values() if str(v).strip())


def _header(document: Document, name: str, contact: dict[str, str] | None) -> None:
    """Render the name/contact block shared by the CV and the cover letter."""
    if name:
        _heading(document, name, level=0)
    line = _contact_line(contact)
    if line:
        document.add_paragraph(line)


def _date_range(item: Experience) -> str:
    """Render "start – end | location" from whichever parts are present."""
    dates = " – ".join(p for p in (item.start_date, item.end_date) if p)
    return " | ".join(p for p in (dates, item.location) if p)


def _render_experience(document: Document, experience: Experience) -> None:
    heading = " — ".join(p for p in (experience.title, experience.company) if p)
    paragraph = document.add_paragraph()
    paragraph.add_run(heading).bold = True
    dates = _date_range(experience)
    if dates:
        document.add_paragraph(dates)
    for bullet in experience.bullets:
        _bullet(document, bullet)


def _render_project(document: Document, project: Project) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run(project.name).bold = True
    if project.description:
        document.add_paragraph(project.description)
    if project.technologies:
        document.add_paragraph(f"Technologies: {', '.join(project.technologies)}")
    if project.url:
        document.add_paragraph(project.url)


def render_cv(
    cv: TailoredCV,
    out_path: Path,
    name: str = "",
    contact: dict[str, str] | None = None,
) -> Path:
    """Render a `TailoredCV` to a .docx file.

    Section order is fixed and mirrors the schema: name/contact header,
    headline, summary, experiences, projects, skills. `relevance_notes` is
    internal reasoning and is deliberately **not** rendered.

    Args:
        cv: The validated tailored CV.
        out_path: Destination .docx path; parent directories are created.
        name: Candidate name (from the `CareerProfile`; `TailoredCV` has none).
        contact: Candidate contact fields, rendered in profile order.

    Returns:
        The path written.
    """
    document = _new_document()
    _header(document, name, contact)
    if cv.headline:
        document.add_paragraph(cv.headline)
    if cv.summary:
        _heading(document, "Summary", level=1)
        document.add_paragraph(cv.summary)
    if cv.selected_experiences:
        _heading(document, "Experience", level=1)
        for experience in cv.selected_experiences:
            _render_experience(document, experience)
    if cv.selected_projects:
        _heading(document, "Projects", level=1)
        for project in cv.selected_projects:
            _render_project(document, project)
    if cv.highlighted_skills:
        _heading(document, "Skills", level=1)
        document.add_paragraph(", ".join(cv.highlighted_skills))

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out_path))
    logger.debug("docx_renderer: wrote CV %s", out_path)
    return out_path


def render_cover_letter(
    letter: CoverLetter,
    out_path: Path,
    name: str = "",
    contact: dict[str, str] | None = None,
) -> Path:
    """Render a `CoverLetter` to a .docx file (greeting, body, closing).

    Args:
        letter: The generated cover letter.
        out_path: Destination .docx path; parent directories are created.
        name: Candidate name, rendered in the header and under the closing.
        contact: Candidate contact fields, rendered in profile order.

    Returns:
        The path written.
    """
    document = _new_document()
    _header(document, name, contact)
    if letter.greeting:
        document.add_paragraph(letter.greeting)
    for paragraph_text in letter.body_paragraphs:
        document.add_paragraph(paragraph_text)
    if letter.closing:
        document.add_paragraph(letter.closing)
    if name:
        document.add_paragraph(name)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out_path))
    logger.debug("docx_renderer: wrote cover letter %s", out_path)
    return out_path


def convert_to_pdf(docx_path: Path) -> Path | None:
    """Convert a rendered .docx to PDF with headless LibreOffice.

    Args:
        docx_path: The .docx to convert; the PDF lands beside it.

    Returns:
        The PDF path, or ``None`` when conversion is disabled (`RENDER_PDF`),
        LibreOffice is not installed, or the conversion failed — the caller
        keeps the .docx either way.
    """
    if not config.RENDER_PDF:
        return None
    docx_path = Path(docx_path)
    with tempfile.TemporaryDirectory() as profile_dir:
        # LibreOffice needs a writable user profile; HOME is not reliably
        # writable in a container, so point it at a throwaway directory.
        command = [
            config.LIBREOFFICE_BIN,
            f"-env:UserInstallation=file://{profile_dir}",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(docx_path.parent),
            str(docx_path),
        ]
        try:
            subprocess.run(
                command,
                check=True,
                capture_output=True,
                timeout=config.LIBREOFFICE_TIMEOUT_S,
            )
        except FileNotFoundError:
            logger.warning(
                "docx_renderer: %s not found — skipping PDF for %s",
                config.LIBREOFFICE_BIN,
                docx_path.name,
            )
            return None
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
            logger.warning(
                "docx_renderer: PDF conversion of %s failed (%s)", docx_path.name, exc
            )
            return None

    pdf_path = docx_path.with_suffix(".pdf")
    if not pdf_path.exists():
        logger.warning(
            "docx_renderer: LibreOffice reported success but %s is missing", pdf_path
        )
        return None
    return pdf_path
