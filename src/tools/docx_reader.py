"""Deterministic .docx CV reader — headers become section boundaries."""

from pathlib import Path

from docx import Document

from src.models.schemas import SourceDocument


def read_docx(path: str | Path) -> SourceDocument:
    """Extract text from a .docx file, marking heading paragraphs as sections.

    Args:
        path: Path to the .docx file.

    Returns:
        A SourceDocument with heading-delimited raw text.
    """
    path = Path(path)
    doc = Document(str(path))
    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "") if para.style else ""
        if style.startswith("Heading") or style == "Title":
            lines.append(f"\n## {text}")
        else:
            lines.append(text)
    return SourceDocument(
        id=f"cv_docx:{path.name}",
        source_type="cv_docx",
        raw_text="\n".join(lines).strip(),
    )
